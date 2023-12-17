import json
from docx import Document


def extract_chessly_lines(doc_path):
    doc = Document(doc_path)
    data = []
    current_section = None
    current_line = None

    for para in doc.paragraphs:
        if para.text.strip() and not para.text.startswith('Chessly Line'):
            # Starting a new section
            if current_section is not None:
                data.append(current_section)
            current_section = {"title": para.text.strip(), "lines": []}
            current_line = None
        elif para.text.startswith('Chessly Line'):
            # Starting a new line within the current section
            current_line = {"name": para.text.strip(), "moves": []}
            current_section["lines"].append(current_line)

            # Find the next table after this paragraph for move data
            element_index = doc.element.body.index(para._element)
            for next_elem in doc.element.body[element_index + 1:]:
                if next_elem.tag.endswith('tbl'):
                    table = [
                        tbl for tbl in doc.tables if tbl._element is next_elem
                    ][0]
                    for row in table.rows:
                        cells = row.cells
                        # avoid the header row
                        if cells[0].text == 'Move Number':
                            continue
                        if len(cells) >= 3:
                            move = {
                                "Move Number": cells[0].text,
                                "White": cells[1].text.strip(),
                                "Black": cells[2].text.strip()
                            }
                            current_line["moves"].append(move)
                    break

    if current_section is not None:
        data.append(current_section)

    return data


def save_to_json(data, filename):
    with open(filename, 'w') as json_file:
        json.dump(data, json_file, indent=4)


doc_path = 'files/advance_caro.docx'  # Replace with the word document to use
chess_data = extract_chessly_lines(doc_path)
save_to_json(chess_data, 'advance_caro_chessly_lines.json')
